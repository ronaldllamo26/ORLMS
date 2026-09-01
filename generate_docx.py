import sys
import subprocess

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set default font to Times New Roman, 12pt
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def add_centered_bold_paragraph(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p

def add_centered_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    return p

# --- Title Page ---
# Adding spacing to center vertically roughly
for _ in range(3):
    doc.add_paragraph()

add_centered_bold_paragraph(doc, "DESIGN AND DEVELOPMENT OF AN INTELLIGENT LEGISLATIVE MANAGEMENT SYSTEM WITH GROQ LLM-POWERED SIMILARITY CHECKS, AUTOMATED COMPLETENESS VERIFICATION, AND PLAIN-LANGUAGE PUBLIC REGISTRY FOR THE CITY OF SAN JOSE DEL MONTE, BULACAN")

for _ in range(4):
    doc.add_paragraph()

add_centered_paragraph(doc, "A Capstone")
add_centered_paragraph(doc, "Presented to the Faculty of")
add_centered_paragraph(doc, "The College of Computing Studies")
add_centered_paragraph(doc, "Bestlink College of the Philippines")

for _ in range(4):
    doc.add_paragraph()

add_centered_paragraph(doc, "In Partial Fulfillment")
add_centered_paragraph(doc, "of the Requirements for the Degree")
add_centered_paragraph(doc, "Bachelor of Science in Information Technology")

for _ in range(4):
    doc.add_paragraph()

add_centered_bold_paragraph(doc, "CHRISTIAN FERNAND A. LUGTU")
add_centered_bold_paragraph(doc, "RONALD E. LLAMO")
add_centered_bold_paragraph(doc, "KENNETH B. ADVINCULA")
add_centered_bold_paragraph(doc, "SHEENA MAE-ANN L. ROTO")
add_centered_bold_paragraph(doc, "JULIANA G. GOPOLE")

for _ in range(4):
    doc.add_paragraph()

add_centered_paragraph(doc, "February 2026")

doc.add_page_break()

add_centered_bold_paragraph(doc, "TABLE OF CONTENTS")

doc.add_page_break()

# --- Chapter 1 ---

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(14)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(12)
        if level == 3:
            run.italic = True
    return p

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

add_heading(doc, "PART I – INTRODUCTION", level=1)

# 1.1 Project Background and Motivation
add_heading(doc, "1.1 Project Background and Motivation", level=2)
add_paragraph(doc, "In the constitutional and administrative framework of the Philippines, Local Government Units (LGUs) are granted fiscal and administrative autonomy under Republic Act No. 7160, also known as the Local Government Code of 1991. This autonomy empowers local legislative councils—specifically the Sangguniang Panlungsod in component cities—to enact ordinances and resolutions that address the specific needs of their localities. These local policies serve as the legal backbone for municipal operations, regulating critical sectors such as public health, business licensing, environmental protection, zoning, and local taxation. Because these measures carry the weight of law within the municipality, their drafting, deliberation, voting, and archiving require meticulous speed, structural correctness, and cross-reference validation. Any delay or error in this legislative pipeline directly hinders municipal growth, public safety, and executive operations.")
add_paragraph(doc, "The City of San Jose del Monte (CSJDM) in the province of Bulacan presents a unique and compelling case for legislative modernization. As one of the largest and most populous component cities in Central Luzon, CSJDM is experiencing rapid urbanization, massive housing developments, and an influx of commercial enterprises due to its strategic proximity to Metro Manila. With its population now exceeding half a million residents, the administrative and legislative demands placed upon the Office of the Sangguniang Panlungsod have grown exponentially. The city council must continuously draft, review, and pass new ordinances and resolutions to manage municipal expansion. However, this massive volume of documents has severely strained the traditional, manual workflow of the legislative secretariat, exposing vulnerabilities in document tracking and quality control.")
add_paragraph(doc, "Under the current manual framework in CSJDM Bulacan, the legislative process remains heavily paper-bound and fragmented. When a councilor or department head proposes a new legislative measure, the draft undergoes a physical routing process across various committee offices for first reading, review, and reporting. Deliberations, public hearings, and session voting are logged using manual minutes and physical folders. This reliance on physical documents inevitably results in administrative bottlenecks, misplaced files, and delayed committee actions. Furthermore, once an ordinance is enacted, it is stored in physical filing cabinets or basic local folders, making document retrieval highly inefficient when staff or government officials need to reference old regulations.")
add_paragraph(doc, "Beyond the physical storage issue, the secretariat faces a critical challenge regarding policy consistency and legal completeness. Legislative secretaries do not possess an automated system to verify if a newly drafted ordinance conflicts with or duplicates policies enacted in previous years. Manually scanning decades of municipal records is nearly impossible, resulting in a high risk of enacting redundant or contradictory local laws. Additionally, the manual verification of a draft's structural completeness—such as checking for the inclusion of vital components like the enacting clause, penal provisions, severability clause, and effectivity dates—requires hours of manual scrutiny by legal officers before the draft can even be endorsed for a second reading.")
add_paragraph(doc, "To address these systemic inefficiencies, the researchers developed the Ordinance and Resolution Lifecycle Management System (ORLMS). Built on a robust PHP MVC architecture with a PostgreSQL database, this web-based system digitizes the entire legislative lifecycle, transforming document tracking into an automated, paperless workflow. The core innovation of ORLMS lies in its integration with the Groq Cloud API, which leverages state-of-the-art Llama Large Language Models (LLMs) to run intelligent decision-support checks. Upon submission of a draft, the system’s AI validation engine automatically performs a semantic similarity check, comparing the text against all existing database records to detect duplicates or conflicting terms based on context rather than exact keywords. Simultaneously, the AI conducts a structural completeness verification, scoring the document and flagging missing legal components before it is routed for committee review.")
add_paragraph(doc, "Finally, the ORLMS addresses the gap in public accessibility and civic transparency. Traditionally, local ordinances are written in dense, legal terminology that ordinary citizens find difficult to interpret. The system bridges this gap through a guest-accessible Public Registry Portal. When a law is officially enacted, the system utilizes the Groq AI engine to translate complex legal jargon into an understandable, plain-language summary, which is then published alongside the official document. By combining role-based access controls (RBAC) for secure staff operations, digital committee tracking, automated AI compliance checks, and a plain-language citizen registry, this capstone project provides a comprehensive e-governance solution designed to elevate legislative efficiency and transparency in the City of San Jose del Monte, Bulacan.")

# 1.2 Problem Statement
add_heading(doc, "1.2 Problem Statement", level=2)
add_paragraph(doc, "The Office of the Sangguniang Panlungsod of the City of San Jose del Monte, Bulacan, serves as the city’s primary legislative authority. However, the legislative secretariat currently lacks a dedicated, integrated, and secure software system to manage the lifecycle of ordinances and resolutions. In the absence of an automated platform, the office relies on manual paper processes and basic Microsoft Excel spreadsheets to log and monitor files. This administrative set-up results in major issues concerning legislative speed, policy coherence, and civic transparency.")
add_paragraph(doc, "Specifically, the researchers identified the following technical and operational problems:")
add_heading(doc, "General Problem", level=3)
add_paragraph(doc, "The overarching problem is the lack of a centralized, secure, and intelligent legislative management system in the City of San Jose del Monte, Bulacan. The current reliance on manual routing and disconnected Excel sheets results in significant delays in policy deliberation, high risks of files being misplaced, legal redundancies in newly enacted ordinances, and a complete absence of an online public channel to access local city policies.")
add_heading(doc, "Specific Problems", level=3)
add_bullet(doc, "Absence of Automated Semantic Conflict and Redundancy Detection: The use of basic Excel sheets is limited to simple text logging. It does not provide any semantic validation. When a new ordinance is drafted, legislative staff have no automated tool to cross-reference it against existing records to see if the draft overlaps, duplicates, or contradicts ordinances passed in previous years. This forces staff to manually browse files or rely on memory, leading to legislative redundancies.")
add_bullet(doc, "Tedious and Error-Prone Manual Verification of Legal Structural Completeness: Prior to committee routing, proposed drafts must undergo a structural review. Legal officers and legislative staff must manually inspect the draft to ensure the presence of required components (such as the Title, Enacting Clause, Penal Clause, Severability Clause, and Effectivity Date). This manual review takes hours per document and is highly prone to human error, which delays the bill’s first reading.")
add_bullet(doc, "Inefficient Routing, Status Tracking, and Session Voting: Because there is no digital system to manage workflow progression, tracking the status of a document (e.g., whether it is under committee review, endorsed, or tabled) requires physical follow-ups. Additionally, session voting outcomes and committee reports are logged manually on paper, making it difficult to maintain a reliable, real-time tracking of council members' votes.")
add_bullet(doc, "Lack of Public Access Portal and Plain-Language Translation: The City of San Jose del Monte currently has no existing public website or portal for citizens to view active local laws. To read an ordinance, citizens must physically visit the city hall to request copies. Furthermore, ordinances are written in dense, legal terminology that ordinary residents find difficult to interpret, thereby hindering civic awareness and local compliance.")
add_bullet(doc, "Lack of Immutable Audit Trails and Security Vulnerabilities: Managing legislative records through paper logbooks and Excel sheets poses a severe security risk. Excel files can be easily modified, deleted, or corrupted without leaving a tamper-evident audit history. The office lacks a Role-Based Access Control (RBAC) security framework to track which user created, updated, or archived specific policy records, compromising data integrity.")

# 1.3 Context and Scope
add_heading(doc, "1.3 Context and Scope", level=2)

add_heading(doc, "1.3.1 Context", level=3)
add_paragraph(doc, "Local legislative management in the Philippines remains one of the most underserved areas of public sector digitization, despite the growing administrative demands placed on city and municipal councils. The Sangguniang Panlungsod, as the primary legislative body of component cities, is mandated under Republic Act No. 7160 to continuously enact, amend, and archive ordinances and resolutions governing local public affairs. In practice, however, the secretariats of these legislative offices operate without dedicated software systems, relying instead on physical document routing, manual logbooks, and basic spreadsheet tools that are insufficient to manage the volume, complexity, and legal precision required by modern municipal governance. Academic studies and e-governance initiatives have consistently emphasized that local government units lacking integrated digital workflows experience significant bottlenecks in policy deliberation, high risks of document misplacement, and limited public access to enacted legislation — all of which directly undermine administrative efficiency and civic transparency. The City of San Jose del Monte (CSJDM), Bulacan, with its population of over 651,000 and its rapid urbanization, presents a compelling and urgent case for the digitization of its legislative processes through an intelligent, AI-powered management platform.")

add_heading(doc, "1.3.2 Project Scope", level=3)
add_paragraph(doc, "This capstone project focuses on the design, development, and evaluation of the Ordinance and Resolution Lifecycle Management System (ORLMS), a web-based platform tailored specifically for the Office of the Sangguniang Panlungsod of the City of San Jose del Monte, Bulacan. Built on a PHP MVC architecture and a MySQL database, the system serves as the centralized legislative management platform of the office, enabling the secretariat to manage the full legislative lifecycle — from initial draft submission to official enactment and public publication — within one secure digital environment.")

p = doc.add_paragraph("The scope encompasses ten (10) critical functional areas:")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_bullet(doc, "User Authentication and Role-Based Access Control (RBAC): Restricts permissions and system capabilities based on assigned user roles (Super Admin, Legislative Secretariat, Committee Chair/Members, SP Councilors, LCE/Mayor).")
add_bullet(doc, "Ordinance and Resolution Registry Management: Facilitates digital creation, editing, file uploading, indexing, and tracking of draft ordinances and resolutions.")
add_bullet(doc, "AI-Powered Structural Completeness & Semantic Similarity Validation: Integrates the Groq Cloud API (Llama LLMs) to evaluate structural legal compliance (checking title, enacting clause, penal clause, severability, effectivity) and detect semantic conflicts/redundancies against existing legislative records.")
add_bullet(doc, "Committee Review & Deliberation Workflow: Manages digital document routing to committee desks, allowing members to review drafts alongside AI validation reports, write deliberation notes, and issue committee recommendations.")
add_bullet(doc, "Plenary Reading & Session Vote Tracking: Tracks document progression across first, second, and third readings, recording and tallying councilors' votes during legislative sessions.")
add_bullet(doc, "Executive Approval & Veto Routing: Enables electronic routing of passed measures to the Local Chief Executive (City Mayor) for final signature approval or veto recording with logged veto justification messages.")
add_bullet(doc, "Amendments & Version History Management: Maintains an immutable version history of proposed modifications and textual revisions made throughout the legislative lifecycle.")
add_bullet(doc, "Plain-Language Public Registry Portal: Features a guest-accessible public portal publishing AI-generated plain-language summaries of enacted ordinances alongside official downloadable PDFs to promote civic transparency.")
add_bullet(doc, "Implementation Monitoring Module: Enables legislative staff to track and update the post-enactment enforcement status (Pending, Ongoing, Completed, Delayed) of enacted policies.")
add_bullet(doc, "System Audit Logging: Automatically records tamper-evident logs of all system operations (CREATE, READ, UPDATE, DELETE, LOGIN) with timestamps and user details to protect data integrity.")

add_paragraph(doc, "The project scope includes system analysis, architectural design, database modeling, full web application implementation, and empirical quality evaluation. System quality is evaluated using the ISO/IEC 25010 Software Quality Model across five key quality characteristics: Functional Suitability, Performance Efficiency, Security, Reliability, and Usability. Evaluation respondents are drawn from legislative secretariats, committee members, SP councilors, and IT professionals.")

add_heading(doc, "1.3.3 Project Boundaries (Out-of-Scope)", level=3)
add_paragraph(doc, "The boundaries of this project are defined by its exclusive focus on the legislative operations of the Sangguniang Panlungsod of CSJDM. Specifically, the out-of-scope parameters include:")
add_bullet(doc, "Exclusion of General Municipal Operations: The system does not extend to other general LGU operations of the City Government of San Jose del Monte, such as business permit and licensing, municipal payroll, real property tax assessment, human resource management, civil registry, or law enforcement tracking.")
add_bullet(doc, "Exclusion of External Systems Integration: The platform is restricted to the internal organizational environment of the SP Secretariat and does not integrate with external national government databases, Sangguniang Panlalawigan internal networks, or third-party cloud enterprise systems.")
add_bullet(doc, "Decision-Support Role of Artificial Intelligence: The AI Validation Engine operates strictly in an advisory/decision-support capacity. Final legislative authority over bill drafting, committee endorsement, voting, enactment, and vetoing remains strictly with the SP Councilors and the Local Chief Executive.")

# 1.4 Objectives and Goals
add_heading(doc, "1.4 Objectives and Goals", level=2)
add_heading(doc, "1.4.1 General Objective", level=3)
add_paragraph(doc, "The overarching aim of this study is to design and develop the Ordinance and Resolution Lifecycle Management System (ORLMS), an intelligent web-based platform tailored for the City of San Jose del Monte, Bulacan. The system is intended to streamline local governance by integrating AI-powered compliance verification, semantic similarity checks, and a plain-language citizen registry portal, thereby improving legislative workflow efficiency, policy coherence, and public transparency using a PHP MVC architecture and a PostgreSQL database.")
add_heading(doc, "1.4.2 Specific Objectives", level=3)
p = doc.add_paragraph("To realize this goal, the study sets out the following specific objectives:")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_heading(doc, "System Design and Development", level=3)
add_bullet(doc, "Build a responsive, web-based platform that tracks the entire lifecycle of local ordinances and resolutions from initial drafting to official enactment.")
add_bullet(doc, "Ensure that the system architecture is secure, scalable, and user-friendly, utilizing Role-Based Access Control (RBAC) and a robust PostgreSQL database to protect municipal records.")
add_heading(doc, "AI-Powered Structural Completeness & Semantic Checking", level=3)
add_bullet(doc, "Integrate the Groq Cloud API running Llama models to automatically verify the presence of essential legal components (Title, Enacting Clause, Penal Clause, Severability Clause, and Effectivity Date) in legislative drafts prior to committee routing.")
add_bullet(doc, "Implement a semantic similarity checker that cross-references new drafts against archived PostgreSQL records to alert legislative staff of redundant, overlapping, or conflicting policies.")
add_heading(doc, "AI-Generated Summarization & Plain-Language Public Registry", level=3)
add_bullet(doc, "Develop a guest-accessible Public Portal allowing citizens of San Jose del Monte to easily search and retrieve local ordinances and resolutions online.")
add_bullet(doc, "Leverage Llama Large Language Models (LLMs) to automatically translate complex legal terminology and jargon into simple, plain-language summaries published on the portal to enhance civic transparency.")
add_heading(doc, "Committee Deliberation & Voting Management", level=3)
add_bullet(doc, "Create a digital routing module for committee reviews, enabling members to attach deliberation notes and log official recommendations (Endorse, Reject, or Return for Revision).")
add_bullet(doc, "Establish a digital tallying system to track, log, and monitor session voting outcomes of Sangguniang Panlungsod members during document readings.")
add_heading(doc, "System Evaluation Using ISO/IEC 25010 Standards", level=3)
add_bullet(doc, "Functional Suitability — confirm that the system records legislative data, calculates completeness/similarity scores, and manages workflows accurately.")
add_bullet(doc, "Performance Efficiency — test the system’s ability to process semantic database searches and Groq API requests without significant latency.")
add_bullet(doc, "Security — protect municipal data and audit logs from unauthorized access through secure hashing, session management, and RBAC policies.")
add_bullet(doc, "Usability — evaluate how intuitive and accessible the user interfaces are for both city officials, legislative staff, and the general public.")
add_bullet(doc, "Reliability — ensure the system maintains operational stability, database consistency, and high availability under real-world municipal conditions.")

# 1.5 Significance and Relevance
add_heading(doc, "1.5 Significance and Relevance", level=2)
p = doc.add_paragraph("For the Sangguniang Panlungsod of the City of San Jose del Monte, Bulacan")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_heading(doc, "1. User", level=3)
add_paragraph(doc, "The primary users of the system are the Super Administrators, Legislative Staff, Committee Members, and Sangguniang Panlungsod (SP) Council Members. Each role interacts with the platform according to their legislative responsibilities, supported by Role-Based Access Control (RBAC) to ensure secure, appropriate, and monitored system usage.")

add_heading(doc, "1.a City Administrators and Decision-Makers (LCE / Mayor)", level=3)
add_paragraph(doc, "This study introduces a structured policy-management framework that relies on objective data and automated tracking. By applying AI-driven compliance checks, decision-makers at the executive level (LCE / City Mayor) can ensure that the local laws sent for their final signature are structurally complete and legally compliant, minimizing administrative conflicts and improving the efficiency of local governance.")

add_heading(doc, "1.b Legislative Secretariat and Staff", level=3)
add_paragraph(doc, "The proposed system minimizes the manual workload involved in document routing, logbook entry, and physical archive browsing. It eliminates the manual tracking of first, second, and third readings by digitizing the pipeline. It also supports secretariat operations by automating structural compliance and similarity verification, saving hours of work during session preparations.")

add_heading(doc, "1.c Committee Chairpersons and Council Members", level=3)
add_paragraph(doc, "City councilors and committee chairs benefit from real-time access to legislative drafts and AI validation reports. It allows committee members to evaluate draft contents against existing policies immediately, reducing administrative bottlenecks during committee deliberations, ensuring objective review, and providing transparent voting records.")

add_heading(doc, "1.d The General Public / Citizens of CSJDM Bulacan", level=3)
add_paragraph(doc, "Ordinary citizens and local business owners gain direct, online access to enacted municipal ordinances and resolutions without visiting the city hall. Through the plain-language summaries generated by Llama LLMs, the system translates dense legal jargon into easily digestible text, fostering civic awareness, transparency, and law compliance.")

add_heading(doc, "2. Technical", level=3)
add_paragraph(doc, "The technical beneficiaries of this study include IT students, system developers, programmers, database administrators, and IT educators/professionals. They gain practical insights into how Large Language Model (LLM) APIs and relational database technologies can be integrated into civic management systems, making this research relevant beyond the boundaries of CSJDM.")

add_heading(doc, "2.1 IT and Technical Staff", level=3)
add_paragraph(doc, "The research offers a replicable model for integrating the Groq Cloud API and PostgreSQL databases into web-based applications. It provides the city's IT staff and software developers with a technical guide on building responsive, secure, and role-restricted platforms that can handle sensitive legislative data and manage complex document workflows.")

add_heading(doc, "2.2 Other programmers and developers", level=3)
add_paragraph(doc, "This study contributes to the software development community by demonstrating how modular MVC architectures (in PHP) and open-source database engines (PostgreSQL) can be leveraged to solve record-keeping and redundancy challenges in local government units. It showcases best practices in combining generative AI engines with relational databases to yield structured JSON validation reports.")

add_heading(doc, "2.3 IT Teachers and Related Professionals", level=3)
add_paragraph(doc, "Educators and IT professionals can use this study as an industry-aligned instructional reference. It bridges theoretical concepts in API integration, SQL queries, database indexing, and software engineering with a real-world e-governance implementation, providing a valuable teaching reference for academic and professional training contexts.")

add_heading(doc, "Future Research", level=3)
add_paragraph(doc, "Future Researchers. This study establishes a foundation for further exploration of natural language processing (NLP) in municipal legislation. It opens opportunities to investigate advanced integrations such as automated bill drafting, legal context parsing, multi-lingual translation, and the socio-political impacts of AI adoption in public administration.")

# 1.6 Definition of Terms
add_heading(doc, "1.6 Definition of Terms", level=2)
p = doc.add_paragraph("To provide clarity and consistency throughout this study, the following terms are defined as they are used within the context of the capstone project:")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_bullet(doc, "Local Government Unit (LGU) — The local administrative branch of the government (specifically referring to the City Government of San Jose del Monte, Bulacan) that serves as the primary beneficiary and organizational setting of this study.")
add_bullet(doc, "Ordinance and Resolution Lifecycle Management System (ORLMS) — A web-based platform developed in this research to digitize, manage, validate, and track local ordinances and resolutions through their entire legislative lifecycle.")
add_bullet(doc, "Ordinance — A local municipal law enacted by the local legislative council (Sangguniang Panlungsod) that carries a permanent and general effect on the community and requires LCE approval.")
add_bullet(doc, "Resolution — A formal expression of opinion, statement of policy, or temporary decision passed by the local council that addresses specific issues or administrative matters.")
add_bullet(doc, "Artificial Intelligence (AI) — The simulation of human intelligence processes by computer systems, including text understanding and summarization. In this study, AI is applied to automate semantic similarity checking and plain-language summary generation.")
add_bullet(doc, "Large Language Model (LLM) — A type of deep learning artificial intelligence (specifically referring to the Llama models processed via the Groq Cloud API) trained to understand, validate, and generate human-like legal and plain-language texts.")
add_bullet(doc, "Semantic Similarity Checking — The AI-powered process of evaluating the context, meaning, and intent of two different legislative texts to identify overlaps, redundancies, or policy conflicts, regardless of exact keyword matches.")
add_bullet(doc, "Structural Completeness Verification — An automated compliance check performed by the system's AI engine to score a draft ordinance and verify the presence of mandatory legal components (such as the Title, Enacting Clause, Penal Clause, and Effectivity Date).")
add_bullet(doc, "Role-Based Access Control (RBAC) — A security mechanism that restricts system access based on user roles (Super Admin, Legislative Staff, Committee Member, and SP Member) to protect data confidentiality, integrity, and operational flow.")
add_bullet(doc, "Plain-Language Summary — An AI-generated translation of dense, complex legal terminology into simple, easily understandable words, published on the guest registry to make laws accessible to ordinary citizens.")
add_bullet(doc, "PostgreSQL Database — An open-source relational database management system used in this project to store, index, and manage legislative records, metadata, user details, and system audit logs securely.")
add_bullet(doc, "Public Registry Portal — A guest-accessible web interface that allows the citizens of San Jose del Monte, Bulacan, to search, retrieve, download, and read plain-language summaries of enacted ordinances and resolutions.")
add_bullet(doc, "Implementation Monitoring Logs — A system feature that allows legislative staff to log, track, and update the execution status of enacted policies (e.g., Pending, Ongoing, Completed, or Delayed).")
add_bullet(doc, "ISO/IEC 25010 Standards — An international framework for evaluating software quality. In this study, it is used to assess the functional suitability, performance efficiency, security, reliability, and usability of the developed system.")

# 1.7 Structure of the Document
add_heading(doc, "1.7 Structure of the Document", level=2)
add_paragraph(doc, "This capstone project is organized into five chapters, each addressing a specific aspect of the research and system development:")
add_heading(doc, "Chapter 1 – Introduction", level=3)
add_paragraph(doc, "Provides the background of the study, the context and scope, the problem statement, objectives and goals, the significance and relevance of the research, and the overall structure of the document.")
add_heading(doc, "Chapter 2 – Review of Related Literature and Studies", level=3)
add_paragraph(doc, "Presents existing theories, frameworks, and related works that form the foundation of the proposed system. It highlights current gaps in local legislative processes, policy redundancy tracking, and document validation practices, and explains how this study contributes to addressing them.")
add_heading(doc, "Chapter 3 – Methodology", level=3)
add_paragraph(doc, "Describes the research design, development approach, and tools used in creating the system. It outlines the System Development Life Cycle (SDLC), data collection methods, database structure, and evaluation procedures.")
add_heading(doc, "Chapter 4 – Results and Discussion", level=3)
add_paragraph(doc, "Showcases the developed system, its features, and evaluation results. This chapter analyzes how the system meets the stated objectives and discusses its effectiveness in improving the efficiency, accuracy, and public transparency of ordinance and resolution management.")
add_heading(doc, "Chapter 5 – Summary, Conclusions, and Recommendations", level=3)
add_paragraph(doc, "Summarizes the findings of the study, presents conclusions based on the results, and provides recommendations for future improvements, implementation strategies, and further research in e-governance systems.")

output_path = r'c:\xampp\htdocs\Capstone\docs\documents\Chapter_1_Final_Format.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
